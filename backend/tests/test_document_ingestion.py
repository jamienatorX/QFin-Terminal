import io
import unittest
from unittest.mock import AsyncMock, patch

import pandas as pd
from docx import Document
from fastapi.testclient import TestClient

import main
from document_ingestion import DocumentParseError, MAX_UPLOAD_BYTES, parse_document_bytes


class DocumentIngestionTests(unittest.TestCase):
    def test_csv_extracts_financial_table(self):
        parsed = parse_document_bytes(
            "income.csv",
            "text/csv",
            b"period,revenue,net_income\n2024,1000,100\n2025,1200,140\n",
        )
        self.assertEqual(parsed["kind"], "spreadsheet")
        self.assertEqual(parsed["rows"], 2)
        self.assertIn("revenue", parsed["text"])
        self.assertIn("2025", parsed["text"])
        self.assertEqual(parsed["table_data"][0]["records"][1]["net_income"], 140)

    def test_xlsx_extracts_multiple_sheets(self):
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            pd.DataFrame({"metric": ["Revenue"], "value": [1200]}).to_excel(
                writer, sheet_name="Income", index=False
            )
            pd.DataFrame({"metric": ["Cash"], "value": [400]}).to_excel(
                writer, sheet_name="Balance", index=False
            )

        parsed = parse_document_bytes(
            "annual-report.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            buffer.getvalue(),
        )
        self.assertEqual(parsed["sheets"], ["Income", "Balance"])
        self.assertIn("Revenue", parsed["text"])
        self.assertIn("Cash", parsed["text"])

    def test_docx_extracts_paragraphs_and_tables(self):
        document = Document()
        document.add_heading("Annual Report 2025", level=1)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Revenue"
        table.cell(0, 1).text = "1200"
        table.cell(1, 0).text = "Net income"
        table.cell(1, 1).text = "140"
        buffer = io.BytesIO()
        document.save(buffer)

        parsed = parse_document_bytes(
            "annual-report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
        )
        self.assertEqual(parsed["kind"], "document")
        self.assertIn("Annual Report 2025", parsed["text"])
        self.assertIn("Net income | 140", parsed["text"])

    def test_image_is_encoded_for_vision_without_persistence(self):
        parsed = parse_document_bytes("chart.png", "image/png", b"\x89PNG\r\n\x1a\nimage")
        self.assertEqual(parsed["kind"], "image")
        self.assertTrue(parsed["image_data_url"].startswith("data:image/png;base64,"))
        self.assertNotIn("path", parsed)

    def test_rejects_unsupported_and_oversized_files(self):
        with self.assertRaises(DocumentParseError):
            parse_document_bytes("legacy.doc", "application/msword", b"legacy")
        with self.assertRaises(DocumentParseError):
            parse_document_bytes("too-large.csv", "text/csv", b"x" * (MAX_UPLOAD_BYTES + 1))


class AttachmentEndpointTests(unittest.TestCase):
    def test_multipart_upload_calls_attachment_chat(self):
        expected = {
            "route": {"kind": "finance_concept"},
            "content": "Attachment analysis complete.",
            "facts": None,
            "used_live_data": False,
            "attachment": {"filename": "report.csv"},
            "evidence": {},
            "risk_review": {},
        }
        with patch("main.generate_attachment_reply", new=AsyncMock(return_value=expected)) as generate:
            response = TestClient(main.app).post(
                "/agent/chat/upload",
                data={"message": "Analyze this annual report"},
                files={"file": ("report.csv", b"metric,value\nRevenue,100\n", "text/csv")},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["content"], "Attachment analysis complete.")
        attachment = generate.await_args.args[1]
        self.assertEqual(attachment["kind"], "spreadsheet")
        self.assertIn("Revenue", attachment["text"])


class AttachmentPromptTests(unittest.IsolatedAsyncioTestCase):
    async def test_generic_attachment_prompt_does_not_invent_ticker(self):
        attachment = parse_document_bytes(
            "report.csv",
            "text/csv",
            b"period,revenue,net_income\n2024,1000,100\n2025,1200,140\n",
        )
        with (
            patch("main.ask_qwen", new=AsyncMock(side_effect=main.QwenClientError("timeout"))),
            patch("main.remember_agent_session"),
            patch("main.yahoo_symbol_search", return_value={"symbol": "TSCM"}),
        ):
            result = await main.generate_attachment_reply("Analyze the attached annual report", attachment)

        self.assertEqual(result["route"]["kind"], "document_analysis")
        self.assertIn("Revenue increased", result["content"])
        self.assertIn("+20.0%", result["content"])

    async def test_image_attachment_builds_multimodal_message(self):
        attachment = parse_document_bytes("chart.png", "image/png", b"\x89PNG\r\n\x1a\nimage")
        with (
            patch("main.ask_qwen", new=AsyncMock(return_value="The chart shows improving margins.")) as ask,
            patch("main.remember_agent_session"),
        ):
            result = await main.generate_attachment_reply("Analyze this chart", attachment)

        self.assertEqual(result["content"], "The chart shows improving margins.")
        messages = ask.await_args.args[0]
        user_content = messages[1]["content"]
        self.assertIsInstance(user_content, list)
        self.assertEqual(user_content[1]["type"], "image_url")
        self.assertTrue(user_content[1]["image_url"]["url"].startswith("data:image/png;base64,"))


if __name__ == "__main__":
    unittest.main()
