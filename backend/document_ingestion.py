import base64
import csv
import io
import re
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
from docx import Document
from pypdf import PdfReader


MAX_UPLOAD_BYTES = 15 * 1024 * 1024
MAX_CONTEXT_CHARS = 60_000
MAX_TABLE_ROWS = 250
MAX_TABLE_COLUMNS = 40

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
TEXT_EXTENSIONS = {".txt", ".md", ".rtf"}
SPREADSHEET_EXTENSIONS = {".csv", ".xls", ".xlsx"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_EXTENSIONS = IMAGE_EXTENSIONS | TEXT_EXTENSIONS | SPREADSHEET_EXTENSIONS | DOCUMENT_EXTENSIONS


class DocumentParseError(ValueError):
    pass


def normalize_extracted_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_CONTEXT_CHARS]


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("The text encoding could not be read.")


def strip_rtf(text: str) -> str:
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    return re.sub(r"[{}]", " ", text)


def parse_pdf(data: bytes) -> Dict[str, Any]:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError("The PDF could not be opened.") from exc

    pages: List[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = normalize_extracted_text(page.extract_text() or "")
        except Exception:
            text = ""
        if text:
            pages.append(f"[Page {page_number}]\n{text}")
        if sum(len(item) for item in pages) >= MAX_CONTEXT_CHARS:
            break

    if not pages:
        raise DocumentParseError(
            "No selectable text was found in this PDF. Upload the scanned pages as images so QFin can use vision analysis."
        )
    return {"kind": "pdf", "text": normalize_extracted_text("\n\n".join(pages)), "pages": len(reader.pages)}


def dataframe_to_context(frame: pd.DataFrame, label: str) -> str:
    safe = frame.iloc[:MAX_TABLE_ROWS, :MAX_TABLE_COLUMNS].copy()
    safe.columns = [str(column) for column in safe.columns]
    safe = safe.where(pd.notna(safe), "")
    return f"[{label}]\n{safe.to_csv(index=False)}"


def parse_spreadsheet(data: bytes, extension: str) -> Dict[str, Any]:
    try:
        if extension == ".csv":
            text = decode_text(data)
            try:
                dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
                separator = dialect.delimiter
            except csv.Error:
                separator = ","
            frame = pd.read_csv(io.StringIO(text), sep=separator)
            sheets = {"CSV": frame}
        else:
            sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
    except Exception as exc:
        raise DocumentParseError("The spreadsheet could not be parsed. Check that it is not password-protected or corrupted.") from exc

    contexts = [dataframe_to_context(frame, str(name)) for name, frame in sheets.items()]
    text = normalize_extracted_text("\n\n".join(contexts))
    if not text:
        raise DocumentParseError("The spreadsheet did not contain readable cells.")
    return {
        "kind": "spreadsheet",
        "text": text,
        "sheets": list(sheets.keys()),
        "rows": sum(len(frame.index) for frame in sheets.values()),
    }


def parse_docx(data: bytes) -> Dict[str, Any]:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError("The Word document could not be opened.") from exc

    blocks: List[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table_number, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows[:MAX_TABLE_ROWS]]
        if rows:
            blocks.append(f"[Table {table_number}]\n" + "\n".join(rows))

    text = normalize_extracted_text("\n\n".join(blocks))
    if not text:
        raise DocumentParseError("The Word document did not contain readable text or tables.")
    return {"kind": "document", "text": text, "paragraphs": len(document.paragraphs), "tables": len(document.tables)}


def parse_image(data: bytes, extension: str, content_type: str) -> Dict[str, Any]:
    mime_type = content_type if content_type.startswith("image/") else {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
    }.get(extension, "image/jpeg")
    encoded = base64.b64encode(data).decode("ascii")
    return {"kind": "image", "text": "", "image_data_url": f"data:{mime_type};base64,{encoded}"}


def parse_document_bytes(filename: str, content_type: str, data: bytes) -> Dict[str, Any]:
    safe_name = Path(filename or "upload").name
    extension = Path(safe_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            "Unsupported file type. Use PDF, CSV, XLS, XLSX, DOCX, TXT, Markdown, RTF, PNG, JPG, JPEG, WEBP, or GIF."
        )
    if not data:
        raise DocumentParseError("The uploaded file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentParseError("The uploaded file exceeds the 15 MB limit.")

    if extension == ".pdf":
        parsed = parse_pdf(data)
    elif extension in SPREADSHEET_EXTENSIONS:
        parsed = parse_spreadsheet(data, extension)
    elif extension == ".docx":
        parsed = parse_docx(data)
    elif extension in IMAGE_EXTENSIONS:
        parsed = parse_image(data, extension, content_type or "")
    else:
        text = decode_text(data)
        if extension == ".rtf":
            text = strip_rtf(text)
        parsed = {"kind": "text", "text": normalize_extracted_text(text)}

    parsed.update(
        {
            "filename": safe_name,
            "content_type": content_type or "application/octet-stream",
            "size_bytes": len(data),
        }
    )
    return parsed
